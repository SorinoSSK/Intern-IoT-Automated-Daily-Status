import json
import time
import requests
import datetime
import os, shutil
import pandas as pd
from os.path import exists

from docx import Document
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import ssl

### Global Variables
MADS_FILE_NAME              = "mads_config.xlsx"
VFT_FILE_NAME               = "vft_config.xlsx"
UNITS_SHEET                 = "daily_report_units_sheet"
HOURLY_UNITS_SHEET          = "hourly_report_units_sheet"
DATA_SHEET                  = "data_sheet"
ACCOUNT_SHEET               = "account_sheet"
EMAIL_SHEET                 = "email_sheet"
OUTPUT_FILE                 = "unit_status.docx"
# Account details and fields to be initialized
DATAPLICITY_LOGIN = {}
ACCOUNT_LOGIN               = "to intialize"
LOGS_DISPLAY_PAGE_SIZE      = "to intialize"
LOGS_DISPLAY_PAGE_NUMBER    = "to intialize"
WITHIN_HOURS                = "to intialize"
WITHIN_DAYS                 = "to intialize"
VALIDATION_EMAIL_TIME       = "13:00"
DAILY_EMAIL_TIME            = "14:00"
# Email details and authentication
SMTP_SERVER                 = None
SENDER_EMAIL                = None
RECIPIENTS_EVERYONE         = None
RECIPIENTS_IOT_TEAM         = None
EMAIL_API_KEY               = None
PORT                        = None

# Error message
FAILED_RETRIEVAL = "Likely a server issue. Refresh the unit's logs data page on platform."

#======================= Configuration =======================#
# As of June 2023, MADs configuration is not updated to VFT standard.
# These configuration is called everytime generate_report function is executed.
def configure_mads():
    configure_account_fields(MADS_FILE_NAME)
    configure_data_fields(MADS_FILE_NAME)
    configure_email_fields(MADS_FILE_NAME)
    
    df = pd.read_excel(io=MADS_FILE_NAME, sheet_name=UNITS_SHEET)
    df = df.fillna("")
    df.drop(df.columns[4], axis=1, inplace=True)
    
    units = {}
    for _, row in df.iterrows():
        units[row[0]] = (row[1], row[2], row[3])
    return units

def configure_vft(isHourly=False):
    configure_account_fields(VFT_FILE_NAME)
    configure_data_fields(VFT_FILE_NAME)
    configure_email_fields(VFT_FILE_NAME)
    # There are different sets of email to be sent.
    # Read a different set of unit list if they are hourly
    # Require states that this script will only read all unit at 1pm for validation and 2pm for all listed personnel to save on system resource.
    if isHourly:
        sheet_to_read = HOURLY_UNITS_SHEET
    else:
        sheet_to_read = UNITS_SHEET
    df = pd.read_excel(io=VFT_FILE_NAME, sheet_name=sheet_to_read)
    df = df.fillna("")
    df.drop(df.columns[4], axis=1, inplace=True)
    
    units = {}
    for _, row in df.iterrows():
        units[row[0]] = (row[1], row[2], row[3])
    return units

# These accounts will be used to login into VFlowTechIot.com and dataplicity.com
def configure_account_fields(PLATFORM_FILE_NAME):
    global ACCOUNT_LOGIN
    ACCOUNT_LOGIN = {} # dict where keys are email and password
    
    df = pd.read_excel(io=PLATFORM_FILE_NAME, sheet_name=ACCOUNT_SHEET)
    ACCOUNT_LOGIN["email"] = df.iloc[0][1]
    ACCOUNT_LOGIN["password"] = df.iloc[1][1]
    DATAPLICITY_LOGIN["email"] = df.iloc[2][1]
    DATAPLICITY_LOGIN["password"] = df.iloc[3][1]

def configure_data_fields(PLATFORM_FILE_NAME):
    global LOGS_DISPLAY_PAGE_SIZE, LOGS_DISPLAY_PAGE_NUMBER, WITHIN_HOURS, WITHIN_DAYS
    
    df = pd.read_excel(io=PLATFORM_FILE_NAME, sheet_name=DATA_SHEET)
    WITHIN_HOURS = int(df.iloc[0][1])
    WITHIN_DAYS = int(df.iloc[1][1])
    LOGS_DISPLAY_PAGE_SIZE = int(df.iloc[2][1])
    LOGS_DISPLAY_PAGE_NUMBER = int(df.iloc[3][1])
    
# This function configures the outgoing email specification such as the recipients.
def configure_email_fields(PLATFORM_FILE_NAME):
    def parse_recipients_field(recipients):
        lst = recipients.split(',')
        for i in range(len(lst)):
            lst[i] = lst[i].strip()
        return lst
        
    global SMTP_SERVER, SENDER_EMAIL, RECIPIENTS_IOT_TEAM, RECIPIENTS_EVERYONE, EMAIL_API_KEY, PORT
    df = pd.read_excel(io=PLATFORM_FILE_NAME, sheet_name=EMAIL_SHEET)

    SMTP_SERVER = df.iloc[0][1]
    PORT = df.iloc[1][1]
    SENDER_EMAIL = df.iloc[2][1]
    EMAIL_API_KEY = df.iloc[3][1]
    # Configure recipients for hourly or daily.
    RECIPIENTS_IOT_TEAM = parse_recipients_field(df.iloc[4][1])
    RECIPIENTS_EVERYONE = parse_recipients_field(df.iloc[5][1])
    
### ---------- Logic V1 ----------###
# ONLINE    -- Dataplicity online + platform showing logs in the last WITHIN_HOURS
# PARTIAL   -- Dataplicity online + platform showing logs in the last WITHIN_DAYS
# OFFLINE   -- Dataplicity offline or not showing logs from last WITHIN_DAYS
### ---------- Logic V2 ----------###
# Combination 1
## Dataplicity online + platform showing logs in the last WITHIN_HOURS
### Display Dataplicity ONLINE, VFlowTechIoT ONLINE
# Combination 2
## Dataplicity offline + platform showing logs in the last WITHIN_HOURS
### Display Dataplicity OFFLINE, VFlowTechIoT ONLINE, Remarks to display warning of device disconnection
# Combination 3
## Dataplicity online + platform showing logs in the last WITHIN_DAYS
### Display Dataplicity ONLINE, VFlowTechIoT PARTIAL, Remarks to display warning of data lag
# Combination 4
## Dataplicity offline + platform showing logs in the last WITHIN_DAYS
### Display Dataplicity OFFLINE, VFlowTechIoT PARTIAL, Remarks to display warning of device disconnection and data lag
# Combination 3
## Dataplicity online + not showing logs from last WITHIN_DAYS
### Display Dataplicity ONLINE, VFlowTechIoT OFFLINE, Remarks to display warning of IoT disconnection
# Combination 4
## Dataplicity offline + not showing logs from last WITHIN_DAYS
### Display Dataplicity OFFLINE, VFlowTechIoT OFFLINE

# This function will only read dataplicity function
def run_dataplicity_status():
    # get auth token
    url = "https://apps.dataplicity.com/auth/"
    rq = requests.post(url, data=DATAPLICITY_LOGIN)
    token = rq.json()["token"]

    endpoint = "https://apps.dataplicity.com/devices/"
    headers = {"Authorization": f"Token {token}"}
    try:
        response = requests.get(endpoint, headers=headers)
        json_dump = response.json()
        statuses = {}
        for unit in json_dump:
            # offline by default
            statuses[unit["name"]] = "offline"
            if unit["online"]:
                statuses[unit["name"]] = "online"
        return statuses
    except Exception as ex:
        print(ex)

# This function is no longer up to date as of June 2023
def run_mad_status(dataplicity_status, units):
    os.makedirs("data_dump", exist_ok=True)
    
    # key: name, value: (online/offline, loc, remarks)
    status = {}
    count = 1

    # get auth token
    url = "https://datakrewtech.com/api/sign-in"
    rq = requests.post(url, data=ACCOUNT_LOGIN)
    token = rq.json()["access_token"]

    curr_time = get_current_time()
    start_time = get_partial_from(curr_time, WITHIN_DAYS) # consider logs from WITHIN_DAYS ago

    for key, values in units.items():
        # get details
        unit_name = values[0]
        loc = values[1]
        remarks = values[2] # to display if offline or partial

        # Likely a mismatch in unit's name since all units on platform must be linked to dataplicity
        if unit_name not in dataplicity_status:
            print("No matching name on dataplicity for " + unit_name)
            continue
        
        # Unit is offline on dataplicity
        if dataplicity_status[unit_name] == "offline":
            print("Dataplicity indicates " + unit_name + " is offline")
            status[unit_name] = ("offline", loc, remarks)
            continue
        
        endpoint = (
            "https://datakrewtech.com/api/iot_mgmt/orgs/3/projects/70/gateways/"
            + str(key)
            + "/data_dump_index"
        )
        headers = {"Authorization": f"Bearer {token}"}
        params = {
            "page_size": LOGS_DISPLAY_PAGE_SIZE,
            "page_number": LOGS_DISPLAY_PAGE_NUMBER,
            "to_date": curr_time,
            "from_date": start_time,
        }    

        response = requests.get(endpoint, headers=headers, params=params)

        if response.status_code != 200:
            print("Error in fetching " + unit_name + " data for MADs, HTTP status code: ", response.status_code)
            status[unit_name] = ("error", loc, FAILED_RETRIEVAL)
            if response.status_code == 500:
                response = retry_ping(response, unit_name, endpoint, headers, params, 0)
            continue # do not process further

        try:
            json_dump = response.json()
            json_string = json.dumps(json_dump, indent=4)
            filename = "data_dump/data_dump_" + str(count) + ".json"
            with open(filename, "w") as outfile:
                outfile.write(json_string)
                
        except json.JSONDecodeError:
            print(
                "JSONDecodeError for "
                + unit_name
                + ", check if unit_id is entered correctly in config.json"
            )

        count += 1

        start_track = get_online_from(curr_time, WITHIN_HOURS) # must show data WITHIN_HOURS to be considered 'online'
        data_logs = json_dump["data_dumps"]
        
        
        if len(data_logs) > 0:
            timestamp_epoch = data_logs[0]["data"]["timestamp"]

            if timestamp_epoch * 1000 >= start_track:
                status[unit_name] = ("online", loc, "")
                print(unit_name + ": " + "online")
            else:
                status[unit_name] = ("partial", loc, remarks)
                print(unit_name + ": " + "partial data in the last " + str(WITHIN_DAYS) + " days")

            num_entries = json_dump["total_entries"]
            
        else: # no logs found
            remarks = "No logs data in the last " + str(WITHIN_DAYS) + " days" # overwrite
            status[unit_name] = ("offline", loc, remarks)
            print(unit_name + " found but no logs data in the last " + str(WITHIN_DAYS) + " days")

    return status

# Error occurs (HTML Code: 500), when a gateway is not called for a period of time.
def retry_ping(response, unit_name, endpoint, headers, params, count):
    if count < 5:
        print("Error probably due to delay in data loading, sleep for 5 seconds.")
        time.sleep(5)
        response = requests.get(endpoint, headers=headers, params=params)
        if response.status_code == 500:
            retry_ping(response, unit_name, endpoint, headers, params, count+1)
        else:
            if response.status_code != 200:
                print("Error in fetching " + unit_name + " data, HTTP status code: ", response.status_code)
            return response
    else:
        print("Error in fetching " + unit_name + " data, HTTP status code: ", response.status_code)
        return response

def run_vft_status(dataplicity_status, units):
    os.makedirs("data_dump", exist_ok=True)

    # key: name, value: (Dataplicity online/offline, VFT online/offline, loc, remarks)
    status = {}
    count = 1

    # get auth token
    login_url = "https://backend.vflowtechiot.com/api/sign-in"
    login_rq = requests.post(login_url, data=ACCOUNT_LOGIN)
    login_token = login_rq.json()["access_token"]

    url = "https://backend.vflowtechiot.com/api/orgs/3/sign-in"
    org_headers = {"Auth-Token": f"{login_token}"}
    rq = requests.post(url, headers=org_headers)
    token = rq.json()["access_token"]
    

    curr_time = get_current_time()
    start_time = get_partial_from(curr_time, WITHIN_DAYS) # consider logs from WITHIN_DAYS days ago

    for key, values in units.items():
        # get details
        unit_name = values[0]
        loc = values[1]
        remarks = values[2] # to display if offline or partial

        # Likely a mismatch in unit's name since all units on platform must be linked to dataplicity
        if unit_name not in dataplicity_status:
            print("No matching name on dataplicity for " + unit_name)
            continue
        
        # Logic V1, this code will skip VFT device check if dataplicity is offline.
        # # Unit is offline on dataplicity
        # if dataplicity_status[unit_name] == "offline":
        #     print("Dataplicity indicates " + unit_name + " is offline")
        #     status[unit_name] = ("offline", loc, remarks)
        #     continue
        
        endpoint = (
            "https://backend.vflowtechiot.com/api/iot_mgmt/orgs/3/projects/70/gateways/"
            + str(key)
            + "/data_dump_index"
        )
        headers = {"Authorization": f"Bearer {token}"}
        params = {
            "page_size": LOGS_DISPLAY_PAGE_SIZE,
            "page_number": LOGS_DISPLAY_PAGE_NUMBER,
            "to_date": curr_time,
            "from_date": start_time,
        }

        response = requests.get(endpoint, headers=headers, params=params)

        if response.status_code != 200:
            if dataplicity_status[unit_name] == "offline":
                print(unit_name + " is offline on Dataplicity.")
            else:
                print(unit_name + " is online on Dataplicity.")
            print("Error in fetching " + unit_name + " data for VFT, HTTP status code: ", response.status_code)
            status[unit_name] = ("error", loc, remarks + "\n" + FAILED_RETRIEVAL)
            continue # do not process further

        try:
            json_dump = response.json()
            json_string = json.dumps(json_dump, indent=4)
            filename = "data_dump/data_dump_" + str(count) + ".json"
            with open(filename, "w") as outfile:
                outfile.write(json_string)
                
        except json.JSONDecodeError:
            print(
                "JSONDecodeError for "
                + unit_name
                + ", check if unit_id is entered correctly in config.json"
            )

        count += 1
       
        start_track = get_online_from(curr_time, WITHIN_HOURS) # must show data WITHIN_HOURS to be considered 'online'
        data_logs = json_dump["data_dumps"]
    
        if len(data_logs) > 0:
            timestamp_epoch = data_logs[0]["data"]["timestamp"]

            # platform showing logs in the last WITHIN_HOURS
            if timestamp_epoch * 1000 >= start_track:
                # Dataplicity is offline
                if dataplicity_status[unit_name] == "offline":
                    print(unit_name + " is offline on Dataplicity.")
                    remarks = remarks + "\n" + "Device is disconnected from dataplicity."
                else:
                    print(unit_name + " is online on Dataplicity.")
                status[unit_name] = ("online", loc, remarks)
                print(unit_name + " is online on VFlowTechIoT.")
            else:
                # Dataplicity is offline
                if dataplicity_status[unit_name] == "offline":
                    print(unit_name + " is offline on Dataplicity.")
                    remarks = remarks + "\n" + "Device is disconnected from dataplicity and experiencing data lag?"
                else:
                    print(unit_name + " is online on Dataplicity.")
                    remarks = remarks + "\n" + "Device is experiencing data lag."
                status[unit_name] = ("partial", loc, remarks)
                print(unit_name + ": " + "partial data in the last " + str(WITHIN_DAYS) + " days")

            num_entries = json_dump["total_entries"]
            
        else: # no logs found
            if dataplicity_status[unit_name] == "offline":
                print(unit_name + " is offline on Dataplicity.")
            else:
                print(unit_name + " is online on Dataplicity.")
                remarks = "No logs data in the last " + str(WITHIN_DAYS) + " days. IoT disconnection." # overwrite
            status[unit_name] = ("offline", loc, remarks)
            print(unit_name + " found but no logs data in the last " + str(WITHIN_DAYS) + " days")
    return status
        

### Utils
def _set_cell_background(cell, fill, color = None, val = None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    cell_properties = cell._element.tcPr
    if cell_properties.xpath("w:shd"): # exists existing shading
        cell_shading = cell_properties.xpath("w:shd")[0]
    else: # add new w:shd element
        cell_shading = OxmlElement("w:shd")

    if fill:
        cell_shading.set(qn("w:fill"), fill)
    if color:
        pass #TODO
    if val:
        pass #TODO

    #extend cell props with shading
    cell_properties.append(cell_shading)


def get_current_time():
    return int(time.time() * 1000)

def get_partial_from(curr_time, WITHIN_DAYS):
    return curr_time - 60 * 60 * 24 * WITHIN_DAYS * 1000

def get_online_from(curr_time, WITHIN_HOURS):
    return curr_time - 60 * 60 * WITHIN_HOURS * 1000

def remove_data_dump():
    folder = 'data_dump/'
    for filename in os.listdir(folder):
        if filename != "status.json":
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

def get_greetings(time):
    hour = int(time.split(":")[0])
    if hour >= 4 and hour < 12:
        return "Good Morning,"
    elif hour >= 12 and hour < 18:
        return "Good Afternoon,"
    elif hour >= 18 and hour <= 23:
        return "Good Evening,"
    else: 
        return "Hi,"
    
### Main Code
### Push formatted_time through to get time of execution
### mads: is system running for mads?
### isBlockEmail: should system not send email on report completion?
def generate_report(formatted_time="00:00", mads=False, isBlockEmail=False):
    statusDict = {}
    # rtn = []
    document = Document()
    document.add_heading("Unit Status", 0)
    table = document.add_table(rows=1, cols=5)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "Unit Name"
    if mads:
        heading_cells[1].text = "MADs"
    else:
        heading_cells[1].text = "VFT"
    print("Generating report for", heading_cells[1].text)
    # rtn.append(heading_cells[1].text)
    heading_cells[2].text = "Dataplicity"
    heading_cells[3].text = "Location"                     
    heading_cells[4].text = "Remarks"

    # because there's a dependency between dataplicity status and platform status
    # if dataplicity offline --> unit is offline regardless of logs shown (Logic V1)
    dataplicity_status = None # handled below
    
    if mads:
        tracked_units = configure_mads()
        dataplicity_status = run_dataplicity_status()
        platform_status = run_mad_status(dataplicity_status, tracked_units)
    else:
        tracked_units = configure_vft(isHourly=(formatted_time!=DAILY_EMAIL_TIME and formatted_time!=VALIDATION_EMAIL_TIME))
        dataplicity_status = run_dataplicity_status()
        platform_status = run_vft_status(dataplicity_status, tracked_units)            

    for unit, values in platform_status.items():
        unit_status_platform = values[0]
        unit_status_dataplicity = dataplicity_status[unit]
        location = values[1]
        remark = values[2]
        cells = table.add_row().cells
        if unit_status_platform == "online":
            _set_cell_background(cells[1], "CEEDD0")
        elif unit_status_platform == "offline":
            _set_cell_background(cells[1], fill="F6CACF")
        elif unit_status_platform == "partial":
            _set_cell_background(cells[1], fill="FBEBA6")
        elif unit_status_platform == "error":
            _set_cell_background(cells[1], fill="FF0000")

        if unit_status_dataplicity == "online":
            _set_cell_background(cells[2], "CEEDD0")
        elif unit_status_dataplicity == "offline":
            _set_cell_background(cells[2], fill="F6CACF")
        elif unit_status_platform == "partial":
            _set_cell_background(cells[2], fill="FBEBA6")
        cells[0].text = unit
        cells[1].text = unit_status_platform
        cells[2].text = unit_status_dataplicity
        cells[3].text = location
        cells[4].text = remark
        # print(unit, values)
        statusDict[unit] = values[0]
    
    # rtn.append(statusDict)
    table.style = "Table Grid"

    document.save(OUTPUT_FILE)
    remove_data_dump()                                      # Remove all data dump files
    if not isBlockEmail:                                    # Block Email on initial start up
        if formatted_time == DAILY_EMAIL_TIME or formatted_time == VALIDATION_EMAIL_TIME:   # If it is doing full report
            sendEmail(heading_cells[1].text, formatted_time)
        checkUnitStatus(heading_cells[1].text, statusDict, formatted_time)                  # Check and compare unit status
    StoreStatus(statusDict)                                 # Overwrite status, Always perform status check before overwritting
    # return rtn

# Check status of previous unit
def checkUnitStatus(system, state, formatted_time):
    jsonFile = open('data_dump/status.json')
    unitPrevStatus = json.load(jsonFile)
    offlineUnits = []
    onlineUnits = []
    for i in unitPrevStatus:
        try:
            if i in unitPrevStatus and i in state:
                if unitPrevStatus[i] != state[i] and state[i] == "offline":
                    offlineUnits.append(i)
                elif unitPrevStatus[i] != state[i] and state[i] == "online":
                    onlineUnits.append(i)
        except:
            print("Error Occured")
    # if offlineUnits != []:
    sendEmail(system, formatted_time, isHourly=True, OfflineDevice=offlineUnits, OnlineDevice=onlineUnits)

# Store current status
def StoreStatus(statusDict):
    json_string = json.dumps(statusDict, indent=4)
    filename = "data_dump/status.json"
    with open(filename, "w") as outfile:
        outfile.write(json_string)

# Dynamic sendEmail, System refers to MADs or VFT.
# Formatted time = hh:mm
def sendEmail(System, formatted_time, isHourly=False, OfflineDevice=[], OnlineDevice=[]):
    port = PORT # For SSL
    message = MIMEMultipart()
    
    # If it is an hourly email, send basic,
    # If it is a validation email, send full units read to IoT Team
    # If it is a daily status email, send full units read to Everyone

    if isHourly:
        message['Subject'] = '(Alert) Hourly Status Report'
        message['To'] = ", ".join(RECIPIENTS_IOT_TEAM)                      # Change recipient here
    elif formatted_time == DAILY_EMAIL_TIME:
        message['Subject'] = 'Daily Status Report'
        message['To'] = ", ".join(RECIPIENTS_EVERYONE)                      # Change recipient here
    elif formatted_time == VALIDATION_EMAIL_TIME:
        message['Subject'] = '(Validation) Daily Status Report'
        message['To'] = ", ".join(RECIPIENTS_IOT_TEAM)                      # Change recipient here
        
    message['From'] = SENDER_EMAIL

    if isHourly:                                                            # isHourly is true only when performing check unit status
        body = buildHourlyEmail(System, formatted_time, OfflineDevice, OnlineDevice)
    else:
        body = buildDailyEmail(System, formatted_time)
    message.attach(MIMEText(body, 'plain'))

    with open(OUTPUT_FILE, 'rb') as attachment:
        attachment = MIMEApplication(attachment.read(), _subtype='docx')
        attachment.add_header('content-disposition', 'attachment', filename=OUTPUT_FILE)
        message.attach(attachment)
    
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL(SMTP_SERVER, port, context=context) as server:
        server.login(SENDER_EMAIL, EMAIL_API_KEY)
        server.sendmail(SENDER_EMAIL, message['To'].split(","), message.as_string())

### Craft and send email
def buildDailyEmail(System, formatted_time):
    # add message body
    # old message body
    """ 
    body = get_greetings(formatted_time)
    body += "This Is An Automated Status Report Generated On " + formatted_time + " Hours
    body += "The Status Is Collected From " + System 
    body += "Regards," 
    body += "OfficeUnit-CT3" 
    """

    #new message body
    if formatted_time == VALIDATION_EMAIL_TIME:
        body = "Dear IoT Team,\n"
    elif formatted_time == DAILY_EMAIL_TIME:
        body = "Dear Recipients,\n"
    body += "\n"
    body += "This is an automated status report, generated at " + formatted_time + " hours, regarding the current unit status' from " + System + ".\n"
    body += "\n"
    body += "Please find the report with details attached to this email. \n"
    body += "\n"
    body += "Should you require any further information or have any questions regarding the status report, please do not hesitate to reach out to superadmin@outlook.com.\n"
    body += "\n"
    body += "Thank you for your attention to this matter.\n"
    body += "\n"
    body += "Best Regards,\n"
    body += "VFT IOT Team"
    return body

def buildHourlyEmail(System, formatted_time, OfflineDevice=[], OnlineDevice=[]):
    # add message body
    body = get_greetings(formatted_time) + "\n\n"
    body += "This Is An Hourly Alert Report Generated On " + formatted_time + " Hours\n"
    body += "The Status Is Collected From " + System + "\n\n"
    body += "The Following Device(s) Have Their Status Changed To OFFLINE.\n"
    body += "---------------------------------\n"
    if OfflineDevice != []:
        for i in OfflineDevice:
            body += i + "\n"
    else:
        body += "None\n"
    body += "---------------------------------\n"
    body += "\n"
    body += "The Following Device(s) Have Their Status Changed To ONLINE.\n"
    body += "---------------------------------\n"
    if OnlineDevice != []:
        for i in OnlineDevice:
            body += i + "\n"
    else:
        body += "None\n"
    body += "---------------------------------\n"
    body += "\n"
    body += "Best Regards," + "\n"
    body += "VFT IOT Team" + "\n"
    return body

if __name__ == "__main__":
    print("Starting Script...")
    isMADs = False
    generate_report(mads=isMADs, isBlockEmail=True)
    validSend = True
    while(1):
        try:
            currTime = datetime.datetime.now().strftime("%H:%M")
            # currTime = "17:00"                                                # For debug only

            mins = currTime.split(":")[1]
            if currTime == DAILY_EMAIL_TIME and validSend:
                print("==============================================")
                print("Starting 1400 Status Check.")
                print("==============================================")
                generate_report(formatted_time=currTime, mads=isMADs)           #Generate report
                validSend = False                                               #Flag to stop spam
                print("==============================================")
            elif currTime == VALIDATION_EMAIL_TIME and validSend:
                print("==============================================")
                print("Starting 1300 Status Validation Check.")
                print("==============================================")
                generate_report(formatted_time=currTime, mads=isMADs)           #Generate report
                validSend = False                                               #Flag to stop spam
                print("==============================================")
            elif mins == "30" and not validSend:
                print("==============================================")
                print("Resetting Checks.")
                validSend = True                                                #Reset generate report flag
                print("==============================================")
            elif mins == "00" and validSend:
                print("==============================================")
                print("Starting Hourly Check for", currTime, ".")
                print("==============================================")
                generate_report(formatted_time=currTime, mads=isMADs)           #Generate report
                validSend = False                                               #Flag to stop spam
                print("==============================================")
            time.sleep(30)                                                      #Sleep for 30 seconds
        except Exception as ex:
            print(ex)
